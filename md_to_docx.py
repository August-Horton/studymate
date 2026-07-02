from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import re

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 添加标题样式
    heading_1 = doc.styles.add_style('Heading1', 1)
    heading_1.font.name = '微软雅黑'
    heading_1.font.size = Pt(20)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(0, 51, 102)
    
    heading_2 = doc.styles.add_style('Heading2', 1)
    heading_2.font.name = '微软雅黑'
    heading_2.font.size = Pt(16)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(0, 77, 153)
    
    heading_3 = doc.styles.add_style('Heading3', 1)
    heading_3.font.name = '微软雅黑'
    heading_3.font.size = Pt(14)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(0, 102, 204)
    
    heading_4 = doc.styles.add_style('Heading4', 1)
    heading_4.font.name = '微软雅黑'
    heading_4.font.size = Pt(12)
    heading_4.font.bold = True
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 跳过 YAML front matter
        if i == 0 and line.startswith('---'):
            while i < len(lines) and not lines[i].startswith('---'):
                i += 1
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Heading1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            doc.add_paragraph(line[3:].strip(), style='Heading2')
            i += 1
        elif line.startswith('### '):
            doc.add_paragraph(line[4:].strip(), style='Heading3')
            i += 1
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:].strip(), style='Heading4')
            i += 1
        
        # 表格处理
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
                num_cols = len(headers)
                
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    hdr_cells[j].paragraphs[0].runs[0].bold = True
                    hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                for row_line in table_lines[2:]:
                    row = table.add_row()
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    for j, cell in enumerate(cells):
                        if j < num_cols:
                            row.cells[j].text = cell
                            row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 代码块处理
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1)
        
        # 列表处理
        elif line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                list_items.append(lines[i][2:].strip())
                i += 1
            
            for item in list_items:
                p = doc.add_paragraph()
                p.add_run('● ').font.color.rgb = RGBColor(0, 102, 204)
                p.add_run(item)
        
        # 无序列表（多级）
        elif re.match(r'^\s*[-*] ', line):
            p = doc.add_paragraph()
            level = len(line) - len(line.lstrip())
            p.add_run('● ' * ((level // 2) + 1))
            p.add_run(line.strip()[2:])
        
        # 粗体和斜体处理
        elif line.strip() and not line.startswith('---') and not line.startswith('>'):
            text = line.strip()
            p = doc.add_paragraph()
            
            while text:
                # 处理粗体
                bold_match = re.match(r'\*\*(.*?)\*\*', text)
                italic_match = re.match(r'\*(.*?)\*', text)
                
                if bold_match:
                    p.add_run(bold_match.group(1)).bold = True
                    text = text[len(bold_match.group(0)):]
                elif italic_match:
                    p.add_run(italic_match.group(1)).italic = True
                    text = text[len(italic_match.group(0)):]
                else:
                    p.add_run(text[0])
                    text = text[1:]
        
        else:
            i += 1
    
    doc.save(docx_path)
    print(f"文档已成功保存到: {docx_path}")

if __name__ == '__main__':
    md_file = r'd:\studymate\StudyMate功能说明书.md'
    docx_file = r'd:\studymate\StudyMate功能说明书.docx'
    parse_markdown_to_docx(md_file, docx_file)
